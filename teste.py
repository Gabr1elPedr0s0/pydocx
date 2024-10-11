import os
from docx import Document #biblioteca principal
from docx.shared import Inches, Pt #tamanho e outras coisas relacionadas com fontes
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #alinhamento de paragrafos
from docx.oxml.ns import qn #bordas
from docx.oxml import OxmlElement #bordas

# Função para definir bordas das células
def set_cell_borders(cell):
    """ Define bordas para uma célula """
    cell_element = cell._element
    cell_pr = cell_element.find(qn('w:tcPr'))
    if cell_pr is None:
        cell_pr = OxmlElement('w:tcPr')
        cell_element.insert(0, cell_pr)
    
    cell_borders = cell_pr.find(qn('w:tblCellBorders'))
    if cell_borders is None:
        cell_borders = OxmlElement('w:tblCellBorders')
        cell_pr.append(cell_borders)
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = cell_borders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Tamanho da borda
            cell_borders.append(border)
        else:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')

# Função para criar bilhetes múltiplos com imagem
def cria_bilhete_multiplo(arquivo_nome, imagem_caminho):
    # Ver se a imagem existe
    if not os.path.isfile(imagem_caminho):
        raise FileNotFoundError(f"Imagem não encontrada: {imagem_caminho}")
    
    document = Document()
    
    # Criar uma tabela para por 4 bilhetes por página
    table = document.add_table(rows=2, cols=2)
    table.autofit = False

    for row in table.rows:
        for cell in row.cells:
            # Criar uma tabela interna para alinhar a imagem e o texto
            inner_table = cell.add_table(rows=1, cols=2)
            inner_table.autofit = True
            
            #Largura das colunas internas
            inner_table.columns[0].width = Inches(1.0)  # Imagem
            inner_table.columns[1].width = Inches(3.0)  # Texto
            
            # Adicionar a imagem na primeira coluna da tabela int   erna
            img_cell = inner_table.cell(0, 0)
            img_para = img_cell.add_paragraph()
            img_run = img_para.add_run()
            img_run.add_picture(imagem_caminho, width=Inches(1.0))
            img_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Adicionar o texto na segunda coluna da tabela interna
            text_cell = inner_table.cell(0, 1)
            text_para = text_cell.add_paragraph("Escola SESI de Itapeva")
            text_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            text_para.paragraph_format.space_after = Inches(0.1)  # Ajustar espaçamento

            # Definir o tamanho da fonte
            for run in text_para.runs:
                run.font.size = Pt(10)  # Definir o tamanho da fonte


            # Adicionar um novo parágrafo para o restante do bilhete abaixo da tabela interna
            text_para = cell.add_paragraph(
                "VIAGEM PARA A COPA SESI\n\n" 
                "Querida Família e Estudante\n\n"
                "É com imensa alegria que comunicamos que no dia 26 de Outubro, próxima Sábado, "
                "será realizado os jogos da Copa Sesi na Escola SESI Itapetininga, onde os alunos "
                "aplicarão o protagonismo juvenil através das vivências dos esportes.\n\n"
                "Agradecemos pela atenção e contamos com a sua participação.\n\n"

                "Ciente: ______________________ __/__/____"
            )
            text_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            #Tamanho da fonte
            for run in text_para.runs:
                run.font.size = Pt(10) 

            # Adicionar bordas
            set_cell_borders(cell)

    # Ajustar a largura e Alturas
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(3.5) 
    for row in table.rows:
        row.height = Inches(4.5)  

    #Ajustar as Margens
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Salvar o documento
    document.save(arquivo_nome)

#Chamada
arquivo_nome = "bilhete_multiplo.docx"
imagem_caminho = "logosesi.jfif"  #Imagem

cria_bilhete_multiplo(arquivo_nome, imagem_caminho)
