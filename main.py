from docx import Document
from docx.shared import Inches

def cria_bilhete():

    document = Document()

    document.add_heading('{{bilhete_titulo}}', 0)

    document.add_paragraph('{{bilhete_descricao}}')


    document.add_heading('Coodenadora', level=1)

    document.add_paragraph('{{bilhete_professor}}')


    document.add_heading('Ementa', level=1)



    document.add_paragraph(
        '{{topico_1}}', style='List Bullet'
    )

    document.add_paragraph(
        '{{topico_2}}', style='List Bullet'
    )

    document.add_paragraph(
        '{{topico_3}}', style='List Bullet'
    )

    document.add_paragraph(
        '{{topico_4}}', style='List Bullet'
    )

    document.add_paragraph(
        '{{topico_5}}', style='List Bullet'
    )

    # document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        ('{{data_1}}', '{{topico_1}}'),
        ('{{data_2}}', '{{topico_2}}'),
        ('{{data_3}}', '{{topico_3}}'),
        ('{{data_4}}', '{{topico_4}}'),
        ('{{data_5}}', '{{topico_5}}'),
    
    )

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Data'
    hdr_cells[1].text = 'Topico'

    for data, topico in records:
        row_cells = table.add_row().cells
        row_cells[0].text = data
        row_cells[1].text = topico
        

    document.save('{{arquivo_nome}}')

